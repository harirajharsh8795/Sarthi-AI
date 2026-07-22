import os
import re
import time
import sqlite3
import logging
import urllib.parse
import requests
from tqdm import tqdm
import fitz  # PyMuPDF
import langdetect
from langdetect.lang_detect_exception import LangDetectException
import uuid
import docx
import pytesseract
from PIL import Image, ImageOps, ImageEnhance
import session_manager

# Tesseract executable configuration for Windows
TESSERACT_PATHS = [
    r"C:\Users\HP\AppData\Local\Programs\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files\Tesseract-OCR\tesseract.exe"
]
for path in TESSERACT_PATHS:
    if os.path.exists(path):
        pytesseract.pytesseract.tesseract_cmd = path
        break

# Paths
DATA_DIR = "./data"
SQLITE_DIR = os.path.join(DATA_DIR, "sqlite")
DB_PATH = os.path.join(SQLITE_DIR, "saarthi.db")
CHROMA_DIR = os.path.join(DATA_DIR, "chroma_db")
KB_DIR = os.path.join(DATA_DIR, "knowledge_base")

# Model configuration
EMBEDDING_MODEL_NAME = "sentence-transformers/paraphrase-multilingual-mpnet-base-v2"
OCR_LOW_CONFIDENCE_THRESHOLD = 50.0

# Global lazy-loaded instances
_embedding_model = None
_chroma_client = None

def init_directories():
    """Initializes the required folder structure."""
    domains = ["constitution_and_general_law", "legal", "banking", "hospital"]
    languages = ["en", "hi"]
    
    # Create subfolders for each domain and language
    for d in domains:
        for l in languages:
            os.makedirs(os.path.join(KB_DIR, d, l), exist_ok=True)
            
    os.makedirs(SQLITE_DIR, exist_ok=True)
    os.makedirs(CHROMA_DIR, exist_ok=True)
    os.makedirs("./glossary", exist_ok=True)
    logger.debug("Folder structure initialized successfully.")

import db_manager

def get_db_connection():
    """Returns a connection to the SQLite database using the Singleton pool."""
    return db_manager.db_pool.get_connection()

def init_db():
    """Initializes the SQLite database table."""
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS kb_documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            domain TEXT NOT NULL,
            language TEXT NOT NULL,
            filename TEXT NOT NULL UNIQUE,
            source_url TEXT NOT NULL UNIQUE,
            discovered_via TEXT NOT NULL,
            page_count INTEGER,
            chunk_count INTEGER,
            status TEXT NOT NULL,
            downloaded_at TEXT,
            indexed_at TEXT
        )
    """)
    conn.commit()
    conn.close()
    logger.debug("SQLite database and kb_documents table initialized.")

def get_document_by_url(url):
    """Retrieves a document record from SQLite by source_url."""
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM kb_documents WHERE source_url = ?", (url,))
    row = cursor.fetchone()
    conn.close()
    return dict(row) if row else None

def get_all_documents():
    """Retrieves all document records from SQLite."""
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM kb_documents")
    rows = cursor.fetchall()
    conn.close()
    return [dict(r) for r in rows]

def upsert_document_record(domain, language, filename, source_url, discovered_via, status, 
                           page_count=None, chunk_count=None, downloaded_at=None, indexed_at=None):
    """Inserts or updates a document record in SQLite."""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Check if record exists by source_url or by filename
    cursor.execute("SELECT id FROM kb_documents WHERE source_url = ? OR filename = ?", (source_url, filename))
    row = cursor.fetchone()
    
    now_str = time.strftime("%Y-%m-%d %H:%M:%S")
    
    if row:
        # Update existing record
        doc_id = row['id']
        update_fields = []
        params = []
        
        # Build dynamic update statement based on provided fields
        fields_to_update = {
            "domain": domain,
            "language": language,
            "filename": filename,
            "source_url": source_url,
            "discovered_via": discovered_via,
            "status": status
        }
        if page_count is not None:
            fields_to_update["page_count"] = page_count
        if chunk_count is not None:
            fields_to_update["chunk_count"] = chunk_count
        if downloaded_at is not None:
            fields_to_update["downloaded_at"] = downloaded_at
        if indexed_at is not None:
            fields_to_update["indexed_at"] = indexed_at
            
        for k, v in fields_to_update.items():
            update_fields.append(f"{k} = ?")
            params.append(v)
            
        params.append(doc_id)
        query = f"UPDATE kb_documents SET {', '.join(update_fields)} WHERE id = ?"
        cursor.execute(query, tuple(params))
    else:
        # Insert new record
        cursor.execute("""
            INSERT INTO kb_documents 
            (domain, language, filename, source_url, discovered_via, page_count, chunk_count, status, downloaded_at, indexed_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (domain, language, filename, source_url, discovered_via, page_count, chunk_count, status, downloaded_at, indexed_at))
        
    conn.commit()
    conn.close()

import services

def get_embedding_model():
    """Returns the singleton sentence-transformers model from services."""
    return services.embedding_service.get_model()

def get_chroma_collection():
    """Returns the singleton ChromaDB knowledge_base collection from services."""
    return services.chroma_manager.get_collection("knowledge_base")

def download_file(url, save_path, unstable=False):
    """
    Downloads a file with streamed requests, 30s timeout, and 3 retries with exponential backoff.
    If unstable=True, verifies that the Content-Type is application/pdf before writing to disk.
    """
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    }
    
    max_retries = 3
    backoff_factor = 2
    
    for attempt in range(1, max_retries + 1):
        try:
            logger.debug(f"Downloading {url} (Attempt {attempt}/{max_retries})...")
            response = requests.get(url, headers=headers, stream=True, timeout=30)
            
            # Check response code
            if response.status_code != 200:
                logger.debug(f"Failed with HTTP Status {response.status_code}")
                time.sleep(backoff_factor ** attempt)
                continue
                
            # Content type check
            content_type = response.headers.get("Content-Type", "").lower()
            if "application/pdf" not in content_type:
                # If content-type is missing or unexpected, but it's marked unstable or looks HTML-ish
                if unstable or "text/html" in content_type:
                    logger.debug(f"Warning: URL did not resolve to a PDF. Content-Type is '{content_type}'. Skipping.")
                    return False
            
            # Save the file
            temp_path = save_path + ".tmp"
            with open(temp_path, "wb") as f:
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
                        
            # Rename temp to target
            if os.path.exists(save_path):
                os.remove(save_path)
            os.rename(temp_path, save_path)
            logger.debug(f"Successfully downloaded and saved to {save_path}")
            return True
            
        except Exception as e:
            logger.debug(f"Error on attempt {attempt}: {e}")
            if attempt < max_retries:
                time.sleep(backoff_factor ** attempt)
            else:
                logger.debug("Max retries exceeded.")
                
    return False

def extract_text_and_language(pdf_path, expected_lang):
    """
    Extracts text page-by-page from a PDF using PyMuPDF (fitz).
    Verifies the document language on page 1 (or first readable page) using langdetect.
    Logs warning on mismatch.
    Returns (pages_text, detected_lang).
    """
    doc = fitz.open(pdf_path)
    pages_text = []
    
    # Extract text from all pages
    for i, page in enumerate(doc):
        text = page.get_text()
        pages_text.append((i + 1, text))
        
    doc.close()
    
    # Detect language using first few readable pages
    sample_text = ""
    for _, text in pages_text:
        # Strip spacing and punctuation to see if we have actual word characters
        clean_text = re.sub(r"\s+", " ", text).strip()
        if len(clean_text) > 100:
            sample_text = clean_text[:500]
            break
            
    if not sample_text:
        # Fallback to concatenate all page texts if they are all very short
        sample_text = " ".join([text for _, text in pages_text])
        
    detected_lang = "unknown"
    if sample_text.strip():
        try:
            detected_lang = langdetect.detect(sample_text)
        except LangDetectException:
            pass
            
    # Mismatch check
    if detected_lang != expected_lang and expected_lang in ["hi", "en"]:
        # If detected is 'hi' but expected is 'en' (or vice versa), log warning
        # Sometimes small snippets or numbers make langdetect predict wrong, but let's log warning
        logger.debug(f"Warning: Language mismatch for {os.path.basename(pdf_path)}. Expected '{expected_lang}', detected '{detected_lang}'.")
        
    return pages_text, detected_lang

def chunk_text(pages_text, chunk_size=512, chunk_overlap=80):
    """
    Splits page-by-page text into character-based sliding window chunks of size 512 with 80 character overlap.
    Tracks page number per chunk in a Unicode-safe manner.
    Returns a list of chunk dicts: {"text": str, "page_number": int, "chunk_index": int}
    """
    all_chunks = []
    
    for page_num, text in pages_text:
        # Clean text basic whitespace normalization
        # Note: Do not remove Hindi/Devanagari characters
        normalized_text = re.sub(r"\r\n|\r|\n", "\n", text)
        text_len = len(normalized_text)
        
        if text_len == 0:
            continue
            
        start = 0
        chunk_idx = 0
        
        while start < text_len:
            end = min(start + chunk_size, text_len)
            chunk_content = normalized_text[start:end]
            
            # Save chunk
            all_chunks.append({
                "text": chunk_content,
                "page_number": page_num,
                "chunk_index": chunk_idx
            })
            
            chunk_idx += 1
            start += (chunk_size - chunk_overlap)
            
            # If we reached the end of page and the remaining characters are very few
            if start >= text_len - chunk_overlap and end == text_len:
                break
                
    return all_chunks

def index_document(domain, language, filename, source_url, discovered_via, file_path):
    """
    Orchestrates extraction, language verification, chunking, embedding generation,
    and storing in ChromaDB + SQLite update.
    Returns (page_count, chunk_count, status).
    """
    try:
        logger.debug(f"Processing PDF for indexing: {filename}...")
        
        # 1. Extract text and language
        pages_text, detected_lang = extract_text_and_language(file_path, language)
        page_count = len(pages_text)
        
        if page_count == 0:
            logger.debug(f"Skipping empty or unreadable PDF: {filename}")
            upsert_document_record(
                domain=domain, language=language, filename=filename, source_url=source_url,
                discovered_via=discovered_via, status="failed", page_count=0, chunk_count=0
            )
            return 0, 0, "failed"
            
        # 2. Chunk text
        chunks = chunk_text(pages_text)
        chunk_count = len(chunks)
        logger.debug(f"Extracted {page_count} pages and created {chunk_count} chunks.")
        
        if chunk_count == 0:
            logger.debug(f"No text chunks created for {filename}.")
            upsert_document_record(
                domain=domain, language=language, filename=filename, source_url=source_url,
                discovered_via=discovered_via, status="failed", page_count=page_count, chunk_count=0
            )
            return page_count, 0, "failed"
            
        # 3. Generate embeddings
        model = get_embedding_model()
        texts = [c["text"] for c in chunks]
        embeddings = model.encode(texts, show_progress_bar=False)
        
        # 4. Insert/Upsert into ChromaDB
        collection = get_chroma_collection()
        
        ids = []
        metadatas = []
        import numpy as np
        
        for idx, chunk in enumerate(chunks):
            # Stable deterministic ID
            chunk_id = f"{filename}_p{chunk['page_number']}_c{chunk['chunk_index']}"
            ids.append(chunk_id)
            
            # Calculate embedding norm
            emb_norm = float(np.linalg.norm(embeddings[idx]))
            
            # Metadata structure
            meta = {
                "domain": domain,
                "language": language,
                "filename": filename,
                "source_url": source_url,
                "page_number": chunk["page_number"],
                "chunk_index": chunk["chunk_index"],
                "discovered_via": discovered_via,
                "vector_norm": emb_norm
            }
            metadatas.append(meta)
            
        collection.upsert(
            ids=ids,
            embeddings=[emb.tolist() for emb in embeddings],
            metadatas=metadatas,
            documents=texts
        )
        logger.debug(f"Upserted {chunk_count} chunks into ChromaDB.")
        
        # 5. Update SQLite document record
        now_str = time.strftime("%Y-%m-%d %H:%M:%S")
        upsert_document_record(
            domain=domain, language=language, filename=filename, source_url=source_url,
            discovered_via=discovered_via, status="indexed", page_count=page_count,
            chunk_count=chunk_count, indexed_at=now_str
        )
        return page_count, chunk_count, "indexed"
        
    except Exception as e:
        logger.debug(f"Error indexing document {filename}: {e}")
        upsert_document_record(
            domain=domain, language=language, filename=filename, source_url=source_url,
            discovered_via=discovered_via, status="failed"
        )
        return 0, 0, "failed"

def get_user_docs_collection():
    """Returns the singleton ChromaDB user_docs collection from services."""
    return services.chroma_manager.get_collection("user_docs")

def get_chroma_client():
    """Returns the singleton ChromaDB client from services."""
    return services.chroma_manager.get_client()

def extract_text_from_docx(docx_path):
    """Extracts text from a DOCX file using python-docx."""
    doc = docx.Document(docx_path)
    full_text = []
    
    # Extract from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                text_content = cell.text.strip()
                if text_content and text_content not in row_text:
                    row_text.append(text_content)
            if row_text:
                full_text.append(" | ".join(row_text))
                
    return "\n".join(full_text)

def extract_text_from_image_ocr(image_path):
    """
    Applies robust preprocessing (grayscale, adaptive median filtering for noise, 
    autocontrast stretching, and adaptive thresholding using mean)
    to extract text and compute mean confidence using Tesseract OCR.
    """
    from PIL import ImageFilter
    import numpy as np
    
    img = Image.open(image_path)
    
    # 1. Grayscale
    gray = ImageOps.grayscale(img)
    
    # 2. Adaptive Median Filter to remove salt-and-pepper noise
    denoised = gray.filter(ImageFilter.MedianFilter(size=3))
    diff = np.abs(np.array(gray, dtype=np.int16) - np.array(denoised, dtype=np.int16))
    mean_diff = np.mean(diff)
    processed = denoised if mean_diff > 3.0 else gray
    
    # 3. Autocontrast stretching
    contrast = ImageOps.autocontrast(processed)
    
    # 4. Adaptive thresholding using the mean of the image
    img_np = np.array(contrast)
    mean_val = np.mean(img_np)
    threshold = mean_val if 30 < mean_val < 225 else 127
    binarized = contrast.point(lambda p: 255 if p > threshold else 0)
    
    # Run OCR with english and hindi and return dictionary data for confidence scores
    data = None
    try:
        data = pytesseract.image_to_data(binarized, lang="eng+hin", output_type=pytesseract.Output.DICT)
    except Exception as e:
        logger.debug(f"Warning: OCR with 'eng+hin' failed ({e}). Falling back to 'eng'.")
        try:
            data = pytesseract.image_to_data(binarized, lang="eng", output_type=pytesseract.Output.DICT)
        except Exception as ex:
            logger.debug(f"Error: OCR failed ({ex})")
            
    # Reconstruct text and calculate mean confidence
    text = ""
    mean_conf = 0.0
    if data:
        # Reconstruct lines
        lines = {}
        confidences = []
        for i in range(len(data['text'])):
            word = data['text'][i]
            conf = data['conf'][i]
            if word.strip():
                line_key = (data['page_num'][i], data['block_num'][i], data['par_num'][i], data['line_num'][i])
                if line_key not in lines:
                    lines[line_key] = []
                lines[line_key].append(word)
                
                # Exclude -1 confidences (non-text/block areas)
                if conf > 0:
                    confidences.append(float(conf))
                    
        # Sort line keys and assemble text
        sorted_keys = sorted(lines.keys())
        line_texts = [" ".join(lines[k]) for k in sorted_keys]
        text = "\n".join(line_texts)
        
        if confidences:
            mean_conf = sum(confidences) / len(confidences)
            
    # Fallback to direct OCR on original image if binarization produced sparse/empty text (< 30 chars)
    if len(text.strip()) < 30:
        try:
            raw_text = pytesseract.image_to_string(img, lang="eng+hin")
            if not raw_text.strip():
                raw_text = pytesseract.image_to_string(img, lang="eng")
            if len(raw_text.strip()) > len(text.strip()):
                text = raw_text.strip()
                mean_conf = 75.0
        except Exception as ex:
            logger.debug(f"Direct image OCR fallback failed: {ex}")

    return text, mean_conf

import hashlib
from metrics_manager import metrics_manager
from job_queue import job_queue
from logger_config import logger

def calculate_file_hash(file_path: str) -> str:
    """Computes MD5 hash of a file for duplicate upload checking."""
    hasher = hashlib.md5()
    with open(file_path, 'rb') as f:
        buf = f.read(65536)
        while len(buf) > 0:
            hasher.update(buf)
            buf = f.read(65536)
    return hasher.hexdigest()

def ingest_user_document_task(
    job_id: str, 
    file_path: str, 
    original_filename: str, 
    session_id: str, 
    domain_hint: str = None, 
    conversation_id: str = None, 
    file_hash: str = None, 
    document_id: str = None
):
    """
    Ingest task executed in background thread pool.
    Updates job status/progress and records performance metrics.
    """
    logger = logging.getLogger("saarthi.jobs")
    ext = os.path.splitext(original_filename)[1].lower()
    ocr_used = ext in {".jpg", ".jpeg", ".png", ".webp"}
    
    try:
        # Step 1: Text Extraction
        job_queue.update_job(job_id, "running", 10, "Extracting Text")
        
        pages_text = []
        detected_lang = "unknown"
        ocr_conf = None
        
        start_extract = time.perf_counter()
        if ext == ".pdf":
            pages_text, detected_lang = extract_text_and_language(file_path, "unknown")
            has_text = any(t.strip() for _, t in pages_text)
            
            if not has_text:
                logger.info(f"No text extracted via fitz for {file_path}. Falling back to OCR...")
                ocr_used = True
                
                doc = fitz.open(file_path)
                pages_text = []
                confidences = []
                
                for i, page in enumerate(doc):
                    pix = page.get_pixmap(dpi=150)
                    temp_img_path = f"{file_path}_page_{i}.png"
                    pix.save(temp_img_path)
                    
                    try:
                        text, conf = extract_text_from_image_ocr(temp_img_path)
                        from ocr_sanitizer import ocr_sanitizer
                        text = ocr_sanitizer.sanitize_ocr_text(text)
                        
                        pages_text.append((i + 1, text))
                        if conf > 0:
                            confidences.append(conf)
                    finally:
                        if os.path.exists(temp_img_path):
                            try:
                                os.remove(temp_img_path)
                            except:
                                pass
                
                doc.close()
                if confidences:
                    ocr_conf = sum(confidences) / len(confidences)
                else:
                    ocr_conf = 0.0
                    
        elif ext == ".docx":
            text = extract_text_from_docx(file_path)
            pages_text = [(1, text)]
        else: # Images
            text, ocr_conf = extract_text_from_image_ocr(file_path)
            from ocr_sanitizer import ocr_sanitizer
            text = ocr_sanitizer.sanitize_ocr_text(text)
            pages_text = [(1, text)]
            
        extract_time = time.perf_counter() - start_extract
        if ocr_used:
            metrics_manager.record("ocr_time", extract_time)
            
        page_count = len(pages_text)
        has_text = any(t.strip() for _, t in pages_text)
        if not has_text:
            raise ValueError("No readable text found in document")
            
        if ext in {".docx", ".jpg", ".jpeg", ".png", ".webp"}:
            sample_text = ""
            for _, t in pages_text:
                clean_t = re.sub(r"\s+", " ", t).strip()
                if len(clean_t) > 100:
                    sample_text = clean_t[:500]
                    break
            if not sample_text:
                sample_text = " ".join([t for _, t in pages_text])
            if sample_text.strip():
                try:
                    detected_lang = langdetect.detect(sample_text)
                except Exception:
                    pass

        # Step 2: Chunking
        job_queue.update_job(job_id, "running", 40, "Chunking Text")
        chunks = chunk_text(pages_text)
        chunk_count = len(chunks)
        if chunk_count == 0:
            raise ValueError("Document could not be chunked")

        # Step 3: Embeddings
        job_queue.update_job(job_id, "running", 60, "Generating Embeddings")
        
        start_embed = time.perf_counter()
        # Use singleton embedding service
        model = get_embedding_model()
        texts = [c["text"] for c in chunks]
        embeddings = model.encode(texts, show_progress_bar=False)
        embed_time = time.perf_counter() - start_embed
        metrics_manager.record("embedding_time", embed_time)

        # Step 4: Vector Indexing
        job_queue.update_job(job_id, "running", 80, "Indexing Vectors")
        collection = get_user_docs_collection()
        ids = []
        metadatas = []
        import numpy as np
        
        for idx, chunk in enumerate(chunks):
            chunk_id = f"{session_id}_{document_id}_p{chunk['page_number']}_c{chunk['chunk_index']}"
            ids.append(chunk_id)
            
            # Calculate embedding norm
            emb_norm = float(np.linalg.norm(embeddings[idx]))
            
            meta = {
                "session_id": session_id,
                "conversation_id": conversation_id or "",
                "document_id": document_id,
                "original_filename": original_filename,
                "file_type": ext,
                "page_number": chunk["page_number"],
                "chunk_index": chunk["chunk_index"],
                "language": detected_lang,
                "vector_norm": emb_norm
            }
            if domain_hint:
                meta["domain_hint"] = domain_hint
            metadatas.append(meta)
            
        collection.upsert(
            ids=ids,
            embeddings=[emb.tolist() for emb in embeddings],
            metadatas=metadatas,
            documents=texts
        )
        
        ocr_confidence = ocr_conf if ocr_used else None
        ocr_low_quality_warning = (ocr_confidence < OCR_LOW_CONFIDENCE_THRESHOLD) if ocr_used else False
        
        # Step 5: Finalize status records
        session_manager.update_document_status(
            document_id=document_id,
            session_id=session_id,
            status="indexed",
            page_count=page_count,
            chunk_count=chunk_count,
            ocr_confidence=ocr_confidence,
            ocr_low_quality_warning=ocr_low_quality_warning
        )
        
        job_queue.update_job(job_id, "completed", 100, "Ready")
        
    except Exception as e:
        logger.error(f"Error in background ingestion job {job_id}: {e}")
        try:
            current_ocr_conf = ocr_conf
        except NameError:
            current_ocr_conf = None
            
        ocr_low_quality = (current_ocr_conf < OCR_LOW_CONFIDENCE_THRESHOLD) if current_ocr_conf is not None else False
        
        session_manager.update_document_status(
            document_id=document_id,
            session_id=session_id,
            status="failed",
            ocr_confidence=current_ocr_conf,
            ocr_low_quality_warning=ocr_low_quality
        )
        
        job_queue.update_job(job_id, "failed", 100, "Failed", error_message=str(e))
        
    finally:
        # Guarantee cleanup of temporary uploads
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
                logger.info(f"Purged temporary ingestion file: {file_path}")
            except Exception as cleanup_err:
                logger.error(f"Failed to delete temp file {file_path}: {cleanup_err}")

def ingest_user_document(file_path, original_filename, session_id, domain_hint=None, conversation_id=None):
    """
    Synchronous compat wrapper for legacy caller interfaces.
    Launches and awaits the pipeline synchronously.
    """
    ext = os.path.splitext(original_filename)[1].lower()
    document_id = f"doc_{uuid.uuid4().hex}"
    
    session_manager.create_document_record(
        document_id=document_id,
        session_id=session_id,
        original_filename=original_filename,
        file_type=ext,
        domain_hint=domain_hint,
        ocr_used=ext in {".jpg", ".jpeg", ".png", ".webp"},
        status="pending",
        conversation_id=conversation_id
    )
    
    job_id = job_queue.create_job(document_id, conversation_id)
    ingest_user_document_task(
        job_id, file_path, original_filename, session_id, 
        domain_hint, conversation_id, None, document_id
    )
    
    # Return compatibility format
    doc_record = session_manager.get_document_record(document_id, session_id)
    return {
        "document_id": document_id,
        "ocr_confidence": doc_record.get("ocr_confidence"),
        "ocr_low_quality_warning": doc_record.get("ocr_low_quality_warning"),
        "message": "Ingestion task finished."
    }

def delete_document_and_chunks(document_id, session_id):
    """
    Validates ownership, hard-deletes vectors from ChromaDB,
    and soft-deletes the SQLite metadata row (sets status='deleted').
    """
    # Verify ownership
    doc_record = session_manager.get_document_record(document_id, session_id)
    if not doc_record:
        raise ValueError("Document not found or does not belong to this session")
        
    # Hard-delete chunks from ChromaDB
    collection = get_user_docs_collection()
    collection.delete(where={"document_id": document_id})
    logger.debug(f"Hard-deleted chunks for document {document_id} from ChromaDB.")
    
    # Soft-delete record in SQLite
    session_manager.delete_document_record(document_id, session_id)
    logger.debug(f"Soft-deleted document record {document_id} in SQLite.")
