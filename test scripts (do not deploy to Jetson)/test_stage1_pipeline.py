import os
import shutil
import sqlite3
import uuid
import docx
import fitz
import pytesseract
from PIL import Image, ImageDraw, ImageFont, ImageEnhance, ImageOps
import kb_pipeline
import session_manager

# Ensure Tesseract path is set for this test script too
TESSERACT_PATHS = [
    r"C:\Users\HP\AppData\Local\Programs\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files\Tesseract-OCR\tesseract.exe"
]
for path in TESSERACT_PATHS:
    if os.path.exists(path):
        pytesseract.pytesseract.tesseract_cmd = path
        break

def create_test_pdf(filename, text_content):
    """Helper to create a PDF with PyMuPDF."""
    if os.path.exists(filename):
        os.remove(filename)
    doc = fitz.open()
    page = doc.new_page()
    # Draw simple text
    page.insert_text((50, 50), text_content)
    doc.save(filename)
    doc.close()
    print(f"Created test PDF: {filename}")

def create_test_docx(filename, paragraphs):
    """Helper to create a DOCX file using python-docx."""
    if os.path.exists(filename):
        os.remove(filename)
    doc = docx.Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    # Add a table
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Document Type'
    hdr_cells[1].text = 'Test Agreement'
    doc.save(filename)
    print(f"Created test DOCX: {filename}")

def create_test_image(text, filename, rotate_angle=0, add_noise=False, low_contrast=False):
    """Helper to create text rendered as a PNG image, with optional skew/noise."""
    if os.path.exists(filename):
        os.remove(filename)
    # Create white canvas
    img = Image.new("RGB", (950, 200), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    
    # Try to load a standard Windows font (Arial)
    font_path = r"C:\Windows\Fonts\arial.ttf"
    if os.path.exists(font_path):
        font = ImageFont.truetype(font_path, 28)
    else:
        font = ImageFont.load_default()
        
    draw.text((30, 80), text, fill=(0, 0, 0), font=font)
    
    # Apply skew/rotation
    if rotate_angle != 0:
        img = img.rotate(rotate_angle, fillcolor=(255, 255, 255), resample=Image.BICUBIC)
        
    # Apply low contrast
    if low_contrast:
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(0.15) # Very low contrast
        
    # Apply pixel noise
    if add_noise:
        import random
        random.seed(5)
        pixels = img.load()
        width, height = img.size
        for _ in range(int(width * height * 0.05)): # 5% noise
            x = random.randint(0, width - 1)
            y = random.randint(0, height - 1)
            # Add random color noise
            pixels[x, y] = (random.randint(0, 255), random.randint(0, 255), random.randint(0, 255))
            
    img.save(filename)
    print(f"Created test image: {filename} (rotate={rotate_angle}, noise={add_noise}, low_contrast={low_contrast})")

def calculate_keyword_recovery(text, keywords):
    """Calculates percentage of keywords recovered from the OCR text."""
    if not text:
        return 0.0
    text_lower = text.lower()
    matches = 0
    for kw in keywords:
        if kw.lower() in text_lower:
            matches += 1
    return (matches / len(keywords)) * 100.0

def run_tests():
    print("=================== STARTING STAGE 1 PIPELINE TESTS ===================")
    
    # 1. Initialize environment/databases
    print("Initializing directories and SQLite database...")
    kb_pipeline.init_directories()
    session_manager.init_session_db()
    
    # Generate test files
    temp_dir = "./test_temp"
    os.makedirs(temp_dir, exist_ok=True)
    
    pdf_path = os.path.join(temp_dir, "test_medical.pdf")
    create_test_pdf(pdf_path, "This is a hospital discharge summary. Patient has severe bukhar (fever) and headache.")
    
    docx_path = os.path.join(temp_dir, "test_loan.docx")
    create_test_docx(docx_path, [
        "Personal loan agreement between lender and borrower.",
        "The loan amount is INR 5,00,000 to be repaid in 12 months.",
        "A late fee applies if the payment is delayed."
    ])
    
    # OCR images
    image_clean = os.path.join(temp_dir, "ocr_clean.png")
    image_skewed = os.path.join(temp_dir, "ocr_skewed.png")
    image_noisy = os.path.join(temp_dir, "ocr_noisy.png")
    
    ocr_text = "The patient shows symptoms of high fever and acute bukhar."
    create_test_image(ocr_text, image_clean, rotate_angle=0)
    create_test_image(ocr_text, image_skewed, rotate_angle=7) # Skewed 5-10 degrees
    create_test_image(ocr_text, image_noisy, rotate_angle=0, add_noise=True, low_contrast=True)
    
    # 2. Test Ingestion & Session Isolation
    session_a = f"session_{uuid.uuid4().hex[:8]}"
    session_b = f"session_{uuid.uuid4().hex[:8]}"
    
    print("\n--- Testing Ingestion for Session A ---")
    res_pdf = kb_pipeline.ingest_user_document(pdf_path, "test_medical.pdf", session_a, domain_hint="hospital")
    doc_id_pdf = res_pdf["document_id"]
    res_docx = kb_pipeline.ingest_user_document(docx_path, "test_loan.docx", session_a, domain_hint="banking")
    doc_id_docx = res_docx["document_id"]
    res_ocr_clean = kb_pipeline.ingest_user_document(image_clean, "ocr_clean.png", session_a, domain_hint="hospital")
    doc_id_ocr_clean = res_ocr_clean["document_id"]
    
    print(f"Ingested PDF with ID: {doc_id_pdf}")
    print(f"Ingested DOCX with ID: {doc_id_docx}")
    print(f"Ingested Clean Image with ID: {doc_id_ocr_clean} (Confidence: {res_ocr_clean['ocr_confidence']:.1f}%)")
    
    # Verify metadata in SQLite for Session A
    docs_a = session_manager.get_session_documents(session_a)
    assert len(docs_a) == 3, f"Expected 3 documents for session A, got {len(docs_a)}"
    
    # Verify ChromaDB Isolation
    print("\n--- Testing Session Isolation in ChromaDB ---")
    collection = kb_pipeline.get_user_docs_collection()
    
    # Query with session_a filter
    results_a = collection.get(where={"session_id": session_a})
    print(f"Session A chunks count: {len(results_a['ids'])}")
    assert len(results_a['ids']) > 0, "Session A should have indexed chunks"
    
    # Query with session_b filter
    results_b = collection.get(where={"session_id": session_b})
    print(f"Session B chunks count: {len(results_b['ids'])}")
    assert len(results_b['ids']) == 0, "Session B should return 0 chunks when querying session A's data"
    
    # Verify that all chunks retrieved for session A actually contain session_a metadata
    for meta in results_a['metadatas']:
        assert meta['session_id'] == session_a, f"Leak detected: Chunk contains session {meta['session_id']} instead of {session_a}"
        
    print("Session isolation checks passed successfully.")
    
    # 3. Degraded OCR Testing & Accuracy Evaluation
    print("\n--- Running Degraded OCR Testing ---")
    keywords = ["patient", "symptoms", "high", "fever", "acute", "bukhar"]
    
    # Clean OCR Ingestion
    # We already ingested ocr_clean.png
    clean_rec = session_manager.get_document_record(doc_id_ocr_clean, session_a)
    ocr_clean_text, conf_clean = kb_pipeline.extract_text_from_image_ocr(image_clean)
    accuracy_clean = calculate_keyword_recovery(ocr_clean_text, keywords)
    
    # Skewed OCR Ingestion
    res_ocr_skewed = kb_pipeline.ingest_user_document(image_skewed, "ocr_skewed.png", session_a, domain_hint="hospital")
    doc_id_ocr_skewed = res_ocr_skewed["document_id"]
    ocr_skewed_text, conf_skewed = kb_pipeline.extract_text_from_image_ocr(image_skewed)
    accuracy_skewed = calculate_keyword_recovery(ocr_skewed_text, keywords)
    
    # Noisy OCR Ingestion
    res_ocr_noisy = kb_pipeline.ingest_user_document(image_noisy, "ocr_noisy.png", session_a, domain_hint="hospital")
    doc_id_ocr_noisy = res_ocr_noisy["document_id"]
    ocr_noisy_text, conf_noisy = kb_pipeline.extract_text_from_image_ocr(image_noisy)
    accuracy_noisy = calculate_keyword_recovery(ocr_noisy_text, keywords)
    
    # Assertions for OCR Confidence and Quality Warning
    print("\nVerifying OCR confidence assertions...")
    assert res_ocr_noisy["ocr_low_quality_warning"] is True, "Expected ocr_low_quality_warning == True for noisy image"
    assert res_ocr_noisy["ocr_confidence"] < kb_pipeline.OCR_LOW_CONFIDENCE_THRESHOLD, f"Expected noisy image confidence < {kb_pipeline.OCR_LOW_CONFIDENCE_THRESHOLD}, got {res_ocr_noisy['ocr_confidence']:.1f}%"
    assert res_ocr_clean["ocr_low_quality_warning"] is False, "Expected ocr_low_quality_warning == False for clean image"
    print("OCR confidence assertions passed.")
    
    print("\n=========================================================================")
    print("                      OCR ACCURACY & CONFIDENCE TABLE                    ")
    print("=========================================================================")
    print(f"| Condition              | Accuracy (Keyword Recovery) | Mean Confidence |")
    print(f"|------------------------|-----------------------------|-----------------|")
    print(f"| Clean Image (0 deg)    | {accuracy_clean:22.2f}% | {res_ocr_clean['ocr_confidence']:14.2f}% |")
    print(f"| Skewed Image (7 deg)   | {accuracy_skewed:22.2f}% | {res_ocr_skewed['ocr_confidence']:14.2f}% |")
    print(f"| Noisy/Low-contrast     | {accuracy_noisy:22.2f}% | {res_ocr_noisy['ocr_confidence']:14.2f}% |")
    print("=========================================================================\n")
    
    # 4. Test Idempotency
    print("\n--- Testing Idempotency (Re-uploading) ---")
    # Get old chunks for DOCX
    old_chunks_ids = collection.get(where={"document_id": doc_id_docx})['ids']
    print(f"Old DOCX chunks count: {len(old_chunks_ids)}")
    
    # Re-upload same DOCX file
    res_docx_new = kb_pipeline.ingest_user_document(docx_path, "test_loan.docx", session_a, domain_hint="banking")
    doc_id_docx_new = res_docx_new["document_id"]
    print(f"Re-ingested DOCX with ID: {doc_id_docx_new}")
    
    # Verify old DOCX record in SQLite is soft-deleted
    old_rec = session_manager.get_document_record(doc_id_docx, session_a)
    assert old_rec['status'] == 'deleted', f"Expected old doc to be soft-deleted, but status is {old_rec['status']}"
    
    # Verify old chunks are deleted in ChromaDB
    old_chunks_remaining = collection.get(ids=old_chunks_ids)['ids']
    assert len(old_chunks_remaining) == 0, f"Orphaned chunks from previous version detected: {old_chunks_remaining}"
    
    # Verify new chunks exist
    new_chunks_count = len(collection.get(where={"document_id": doc_id_docx_new})['ids'])
    print(f"New DOCX chunks count: {new_chunks_count}")
    assert new_chunks_count > 0, "New DOCX version should have indexed chunks"
    print("Idempotency checks passed successfully.")
    
    # 5. Test User-Facing Delete
    print("\n--- Testing User-Facing Soft Delete ---")
    # Delete new DOCX document
    kb_pipeline.delete_document_and_chunks(doc_id_docx_new, session_a)
    
    # Verify soft-delete in SQLite
    deleted_rec = session_manager.get_document_record(doc_id_docx_new, session_a)
    assert deleted_rec['status'] == 'deleted', f"Expected document record to be 'deleted', but got {deleted_rec['status']}"
    
    # Verify hard-delete of vector chunks in ChromaDB
    deleted_chunks = collection.get(where={"document_id": doc_id_docx_new})['ids']
    assert len(deleted_chunks) == 0, f"Expected 0 chunks remaining in ChromaDB for deleted document, got {len(deleted_chunks)}"
    print("User-facing delete verification passed successfully.")
    
    # 6. Test File Size Constraints (Try to upload 16MB file)
    print("\n--- Testing Size Validation (Max 15MB) ---")
    large_file_path = os.path.join(temp_dir, "large_dummy.pdf")
    with open(large_file_path, "wb") as f:
        f.write(b"\0" * (16 * 1024 * 1024)) # Write 16MB of dummy bytes
        
    try:
        kb_pipeline.ingest_user_document(large_file_path, "large_dummy.pdf", session_a)
        assert False, "Expected ValueError for size constraint but it succeeded"
    except ValueError as ve:
        print(f"Passed: Correctly rejected large file. Error: {ve}")
        
    # 7. Test File Extension Constraints (Try to upload unsupported extension)
    print("\n--- Testing Extension Validation ---")
    unsupported_file_path = os.path.join(temp_dir, "test.txt")
    with open(unsupported_file_path, "w") as f:
        f.write("Some dummy text")
        
    try:
        kb_pipeline.ingest_user_document(unsupported_file_path, "test.txt", session_a)
        assert False, "Expected ValueError for extension constraint but it succeeded"
    except ValueError as ve:
        print(f"Passed: Correctly rejected unsupported extension. Error: {ve}")
        
    # Cleanup temp test files
    shutil.rmtree(temp_dir)
    print("\n================ STAGE 1 PIPELINE TESTS COMPLETED SUCCESSFULLY ================")

if __name__ == "__main__":
    run_tests()
